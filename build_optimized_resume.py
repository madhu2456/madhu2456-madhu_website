from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "Madhu Dadi_Resume_Optimized.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_borders(cell, color="D9E2EC"):
    tc_pr = cell._tc.get_or_add_tcPr()
    borders = tc_pr.find(qn("w:tcBorders"))
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = f"w:{edge}"
        el = borders.find(qn(tag))
        if el is None:
            el = OxmlElement(tag)
            borders.append(el)
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), "4")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), color)


def set_repeat_table_header(row):
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def style_run(run, bold=False, italic=False, size=None, color=None):
    run.bold = bold
    run.italic = italic
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    run.font.name = "Arial"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")


def set_para(paragraph, before=0, after=0, line_spacing=1.0, alignment=None):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line_spacing
    if alignment is not None:
        paragraph.alignment = alignment


def add_section_heading(doc, title):
    p = doc.add_paragraph()
    set_para(p, before=5, after=2, line_spacing=1.0)
    run = p.add_run(title.upper())
    style_run(run, bold=True, size=8.8, color="1F4D78")
    p_pr = p._p.get_or_add_pPr()
    border = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "5")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "9FB6CF")
    border.append(bottom)
    p_pr.append(border)


def add_bullet(doc, text, level=0):
    p = doc.add_paragraph(style="List Bullet")
    set_para(p, before=0, after=1.15, line_spacing=1.0)
    p.paragraph_format.left_indent = Inches(0.18 + level * 0.18)
    p.paragraph_format.first_line_indent = Inches(-0.12)
    run = p.add_run(text)
    style_run(run, size=7.55, color="111827")
    return p


def add_role(doc, title, organization, dates, location, bullets):
    p = doc.add_paragraph()
    set_para(p, before=2.5, after=0.8, line_spacing=1.0)
    run = p.add_run(title)
    style_run(run, bold=True, size=8.25, color="0B2545")
    run = p.add_run(f" | {organization}")
    style_run(run, bold=True, size=8.25, color="111827")
    run = p.add_run(f" | {dates}")
    style_run(run, size=8.0, color="4B5563")
    if location:
        run = p.add_run(f" | {location}")
        style_run(run, size=8.0, color="4B5563")
    for bullet in bullets:
        add_bullet(doc, bullet)


def add_project(doc, name, role, dates, url, bullets):
    p = doc.add_paragraph()
    set_para(p, before=2.4, after=0.8, line_spacing=1.0)
    run = p.add_run(name)
    style_run(run, bold=True, size=8.2, color="0B2545")
    run = p.add_run(f" | {role}")
    style_run(run, bold=True, size=8.1, color="111827")
    if dates:
        run = p.add_run(f" | {dates}")
        style_run(run, size=7.9, color="4B5563")
    if url:
        run = p.add_run(f" | {url}")
        style_run(run, size=7.9, color="2563EB")
    for bullet in bullets:
        add_bullet(doc, bullet)


def add_skill_table(doc, rows):
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(1.58)
    table.columns[1].width = Inches(5.32)
    for label, detail in rows:
        row = table.add_row()
        row.cells[0].width = Inches(1.58)
        row.cells[1].width = Inches(5.32)
        for cell in row.cells:
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
            set_cell_borders(cell)
            cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        set_cell_shading(row.cells[0], "F1F5F9")
        p = row.cells[0].paragraphs[0]
        r = p.add_run(label)
        style_run(r, bold=True, size=7.35, color="0B2545")
        p = row.cells[1].paragraphs[0]
        r = p.add_run(detail)
        style_run(r, size=7.35, color="111827")
    return table


def add_education_table(doc):
    table = doc.add_table(rows=1, cols=4)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    widths = [Inches(1.1), Inches(1.35), Inches(3.65), Inches(0.95)]
    headers = ["Degree", "Duration", "Institute / Specialization", "Grade"]
    set_repeat_table_header(table.rows[0])
    for i, cell in enumerate(table.rows[0].cells):
        cell.width = widths[i]
        set_cell_shading(cell, "E8EEF5")
        set_cell_borders(cell)
        p = cell.paragraphs[0]
        r = p.add_run(headers[i])
        style_run(r, bold=True, size=7.3, color="0B2545")
    rows = [
        ("MBA", "2018-2020", "Indian Institute of Management, Amritsar | Marketing / Analytics", "3.371/4.33"),
        ("B.Tech.", "2012-2016", "MVGR College of Engineering | Electrical and Electronics Engineering", "79.75%"),
    ]
    for values in rows:
        row = table.add_row()
        for i, value in enumerate(values):
            cell = row.cells[i]
            cell.width = widths[i]
            set_cell_borders(cell)
            p = cell.paragraphs[0]
            r = p.add_run(value)
            style_run(r, size=7.25, color="111827")


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(0.38)
section.bottom_margin = Inches(0.38)
section.left_margin = Inches(0.48)
section.right_margin = Inches(0.48)
section.header_distance = Inches(0.25)
section.footer_distance = Inches(0.25)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "Arial"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
normal.font.size = Pt(7.7)
normal.paragraph_format.space_after = Pt(0)
normal.paragraph_format.line_spacing = 1.0

for style_name in ("Heading 1", "Heading 2", "Heading 3"):
    styles[style_name].font.name = "Arial"
    styles[style_name]._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")

title = doc.add_paragraph()
set_para(title, before=0, after=0, line_spacing=1.0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
run = title.add_run("Madhu Dadi")
style_run(run, bold=True, size=17.2, color="0B2545")

subtitle = doc.add_paragraph()
set_para(subtitle, before=0, after=1.5, line_spacing=1.0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
run = subtitle.add_run("AI & Analytics Leader | LLM/RAG Applications | Healthcare, Marketing & Decision Intelligence")
style_run(run, bold=True, size=8.5, color="1F4D78")

contact = doc.add_paragraph()
set_para(contact, before=0, after=2.2, line_spacing=1.0, alignment=WD_ALIGN_PARAGRAPH.CENTER)
run = contact.add_run(
    "Hyderabad, India | +91 9985422444 | madhu.kumar245@gmail.com | "
    "linkedin.com/in/madhu-dadi-54684531 | madhudadi.in | github.com/madhu2456"
)
style_run(run, size=7.2, color="374151")

add_section_heading(doc, "Summary")
for item in [
    "Senior Analytics Manager with 9+ years of experience driving AI, analytics, and digital transformation across healthcare, travel-tech, media, search marketing, and web/product development.",
    "Builds LLM/RAG applications, agentic workflows, KPI frameworks, Marketing Mix Models, dashboards, and full-stack analytics products that connect business decisions to measurable outcomes.",
    "Proven impact across patient enrollment, digital onboarding, campaign acquisition, churn reduction, global client analytics, and production AI/search visibility products.",
]:
    add_bullet(doc, item)

add_section_heading(doc, "Experience")
add_role(
    doc,
    "Senior Analytics Manager",
    "Novartis",
    "March 2023 - Present",
    "India",
    [
        "Implemented data-driven frameworks across multi-channel digital platforms, driving a 15% increase in patient enrollment.",
        "Architected KPI measurement frameworks that transitioned 90% of patients to digital enrollments and reduced Turnaround Time (TAT) from several days to under 3 hours.",
        "Engineered enterprise-wide patient flows that accelerated digital asset launches from 3 months to 2 weeks.",
        "Won 2 internal AI Hackathons by developing Model Context Protocols (MCPs) and advanced Marketing Mix Models (MMM) for marketing spend optimization.",
        "Directed end-to-end analytics lifecycles, maintaining a 4.9/5 patient satisfaction score through optimized enrollment and adherence communication programs.",
        "Delivered strategic analytical insights to senior leadership and integrated AI-driven marketing strategies, improving operational efficiency by 25%.",
    ],
)
add_role(
    doc,
    "Marketing Analytics Manager",
    "redBus (MMT Group)",
    "June 2022 - March 2023",
    "India",
    [
        "Received Best Performer of the Quarter (Q3 Trailblazer) for a campaign strategy that drove 30% growth in new customer acquisition.",
        "Optimized multi-channel brand campaigns across Facebook, Programmatic, and Google Ads to improve performance and budget efficiency.",
        "Strategized and executed CLM campaigns across Email, WhatsApp, and SMS, reducing customer churn by 20%.",
        "Led marketing analytics measurement for ATL, BTL, and digital campaigns, including impact reporting across redBus markets.",
    ],
)
add_role(
    doc,
    "Analytics Manager",
    "GroupM",
    "September 2020 - May 2022",
    "India",
    [
        "Spearheaded Google Ads Data Hub projects across GroupM agencies, serving the top 5% of global clients.",
        "Received Best Performer of the Year (2020-2021) and earned selection for the WPP NextGen program.",
        "Performed frequency reach analysis to determine optimal ad exposure, improving budget efficiency and campaign margins.",
        "Managed brand lift studies, consumer journey mapping, cluster analysis, and pricing analytics for enterprise client measurement.",
    ],
)
add_role(
    doc,
    "Web Developer and Penetration Tester",
    "Absolinsoft",
    "May 2016 - July 2018",
    "India",
    [
        "Analyzed websites using Google Analytics and executed SEO, SMM, and SEM campaigns, reducing bounce rate by 25%.",
        "Implemented and optimized organic and paid search engine marketing activities, increasing CTR by 20%.",
        "Strategized social media campaigns and documented project/product plans for web applications.",
    ],
)

section = doc.add_section(WD_SECTION.NEW_PAGE)
section.top_margin = Inches(0.38)
section.bottom_margin = Inches(0.38)
section.left_margin = Inches(0.48)
section.right_margin = Inches(0.48)
section.header_distance = Inches(0.25)
section.footer_distance = Inches(0.25)

add_section_heading(doc, "Selected Projects")
add_project(
    doc,
    "Adticks",
    "Full-stack / AI Engineer",
    "2026 - Present",
    "adticks.com",
    [
        "Built an AI visibility and SEO/AEO/GEO auditing platform for large-site diagnostic workflows.",
        "Designed parallel Playwright headless crawls that process 10,000+ pages per audit and compare server HTML with rendered DOM.",
        "Implemented a FastAPI + Celery backend with Redis, PostgreSQL, and 18 parallel audit analyzers for severity-prioritized fix queues.",
        "Audits robots.txt, sitemaps, structured data schemas, and llms.txt configurations, reducing technical audit cycle time by 85%.",
    ],
)
add_project(
    doc,
    "Technical Blog (Python & AI Learning Hub)",
    "Full-stack AI Product",
    "January 2026 - Present",
    "madhudadi.in/blog",
    [
        "Engineered a production-ready technical blog using FastAPI and Next.js, processing a knowledge base of 50+ articles with a RAG system.",
        "Integrated an AI Q&A assistant powered by LangChain and RAG, achieving sub-second response times for complex AI queries.",
        "Optimized the Next.js frontend to 95+ Lighthouse scores for speed and accessibility, contributing to a 20% increase in user engagement.",
    ],
)
add_project(
    doc,
    "Udemy Enroller",
    "Automation Platform",
    "January 2026 - Present",
    "udemyenroller.madhudadi.in",
    [
        "Developed an automated platform that processed 20,000+ free courses within 6 months, saving users an estimated INR 10,00,000+.",
        "Engineered a backend with asynchronous task queues and Redis to handle throughput of 100+ requests per minute.",
        "Implemented custom scrapers to improve data accuracy by 15%, with CAPTCHA handling and rate limiting for resilient automation.",
    ],
)

add_section_heading(doc, "Skills")
add_skill_table(
    doc,
    [
        ("AI / LLM", "Agentic AI, RAG, LangChain, MCPs, Machine Learning, Deep Learning, NLP, MLOps, DevOps"),
        ("Analytics", "KPI frameworks, Marketing Mix Modeling, Hypothesis Testing, A/B Testing, Regression, Clustering, Boosting, Classification, Sentiment Analysis"),
        ("Data / Cloud", "Python, SQL, BigQuery, GCP (Cloud Functions, BigQuery), AWS (S3, Lambda), Big Data, Power BI, Tableau, Looker Studio"),
        ("Marketing", "Performance & Brand Marketing, Loyalty Analytics, Web Analytics, Product Recommendation, CLM, ATL/BTL/Digital measurement"),
        ("Product", "FastAPI, Next.js, PostgreSQL, Redis, Celery, Playwright, Tailwind CSS, GitHub Actions"),
    ],
)

add_section_heading(doc, "Selected Awards")
for item in [
    "Best Performer of the Quarter (Q1 2024), Novartis.",
    "Best Performer of the Quarter (Q3 Trailblazer), redBus, for campaign strategy that resulted in 30% growth in new customers.",
    "Best Performer of the Year (2020-2021), GroupM; selected for the WPP NextGen program.",
    "Special Recognition Award for work as Manager in Aarunya fest, IIM Amritsar.",
]:
    add_bullet(doc, item)

add_section_heading(doc, "Education")
add_education_table(doc)

doc.core_properties.author = "Madhu Dadi"
doc.core_properties.title = "Madhu Dadi Resume"
doc.core_properties.subject = "AI & Analytics Leader Resume"
doc.core_properties.keywords = "AI analytics, RAG, LLM, marketing analytics, healthcare analytics"

doc.save(OUT)
